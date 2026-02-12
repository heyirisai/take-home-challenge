"""
Document text extraction utilities
Supports PDF, DOCX, and TXT files
"""

import os
from typing import Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text from various document formats"""
    
    @staticmethod
    def extract_text(file_path: str) -> Optional[str]:
        """
        Extract text from a document file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text or None if extraction fails
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        extension = Path(file_path).suffix.lower()
        
        try:
            if extension == '.txt':
                return DocumentExtractor._extract_from_txt(file_path)
            elif extension == '.pdf':
                return DocumentExtractor._extract_from_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return DocumentExtractor._extract_from_docx(file_path)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            return ""
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            return ""
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n\n'.join(text_parts)
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> list[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Text to split
            chunk_size: Maximum size of each chunk
            chunk_overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            
            # If not at the end, try to break at a sentence or paragraph
            if end < text_length:
                # Look for paragraph break
                break_point = text.rfind('\n\n', start, end)
                if break_point == -1:
                    # Look for sentence break
                    break_point = text.rfind('. ', start, end)
                    if break_point != -1:
                        break_point += 1
                
                if break_point != -1 and break_point > start:
                    end = break_point
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - chunk_overlap
        
        return chunks
